import os
import tempfile
from datetime import datetime

import streamlit as st
import pandas as pd
import plotly.express as px
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches


# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="📊 Task Completion Analyzer",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------------- STYLING ----------------
st.markdown("""
<style>
section[data-testid="stSidebar"] {
    border-right: 1px solid #e5e7eb;
}
.sidebar-title {
    font-size: 20px !important;
    font-weight: 600 !important;
    text-align: center;
}
.stButton>button {
    border-radius: 10px;
}
</style>
""", unsafe_allow_html=True)

# ---------------- SIDEBAR ----------------
st.sidebar.markdown('<div class="sidebar-title">📂 Task Report Analyzer</div>', unsafe_allow_html=True)

uploaded_files = st.sidebar.file_uploader(
    "📤 Upload Monthly Excel Reports (.xlsx)",
    type=["xlsx"],
    accept_multiple_files=True
)

global_hide_top = st.sidebar.checkbox("🙈 Hide Top Performer", value=False)


# ---------------- CHART UTILITY ----------------
def save_bar_chart(df, x, y, title, path):
    plt.figure(figsize=(8, 4))
    plt.bar(df[x], df[y])
    plt.title(title)
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    plt.savefig(path)
    plt.close()


# ---------------- WORD REPORT GENERATOR ----------------
def generate_combined_word_report(combined_df, monthly_data):
    doc = Document()

    # -------- TITLE PAGE --------
    doc.add_heading("Task Completion Analysis Report", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    # -------- OVERALL SUMMARY --------
    doc.add_heading("Overall Summary", level=1)

    total_completion = int(combined_df["Completion"].sum())
    total_persons = combined_df["Person"].nunique()
    top_person = combined_df.groupby("Person")["Completion"].sum().idxmax()

    doc.add_paragraph(f"Total Completions: {total_completion}")
    doc.add_paragraph(f"Total Active Persons: {total_persons}")
    doc.add_paragraph(f"Top Performer: {top_person}")

    yearly_summary = (
        combined_df.groupby("Person")["Completion"]
        .sum()
        .reset_index()
    )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        save_bar_chart(
            yearly_summary,
            "Person",
            "Completion",
            "Overall Completion per Person",
            tmp.name
        )
        doc.add_picture(tmp.name, width=Inches(6))

    doc.add_page_break()

    # -------- MONTHLY REPORTS --------
    for month, df_month in monthly_data.items():
        doc.add_heading(f"Monthly Report – {month}", level=1)

        doc.add_paragraph(f"Total Completion: {int(df_month['Completion'].sum())}")
        doc.add_paragraph(f"Active Persons: {df_month['Person'].nunique()}")

        # ---- Table ----
        summary = (
            df_month.groupby(["Person", "Portal"])["Completion"]
            .sum()
            .reset_index()
        )

        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Person"
        hdr[1].text = "Portal"
        hdr[2].text = "Completion"

        for _, row in summary.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["Person"])
            cells[1].text = str(row["Portal"])
            cells[2].text = str(int(row["Completion"]))

        # ---- Charts ----
        person_summary = (
            df_month.groupby("Person")["Completion"]
            .sum()
            .reset_index()
        )

        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            save_bar_chart(
                person_summary,
                "Person",
                "Completion",
                f"{month} – Completion per Person",
                tmp.name
            )
            doc.add_picture(tmp.name, width=Inches(6))

        task_summary = (
            df_month.groupby("Task")["Completion"]
            .sum()
            .reset_index()
        )

        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            save_bar_chart(
                task_summary,
                "Task",
                "Completion",
                f"{month} – Task Breakdown",
                tmp.name
            )
            doc.add_picture(tmp.name, width=Inches(6))

        doc.add_page_break()

    file_path = f"Task_Completion_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(file_path)
    return file_path


# ---------------- MAIN LOGIC ----------------
if not uploaded_files:
    st.info("📤 Upload one or more Excel reports to begin analysis.")
    st.stop()

known_tasks = [
    "Preparation and Setup", "Monitor WebInspect", "Quality", "Quality 1", "Quality 2",
    "Authentication and Session", "Access Control", "Input Validation",
    "Business Logic", "Work", "Review", "Remediation 1", "Remediation 2", "Remediation"
]

portals = {
    "AMS PORTAL": [0, 1],
    "EMEA PORTAL": [3, 4],
    "APAC PORTAL": [6, 7],
    "SGP PORTAL": [9, 10],
}

all_data = []
monthly_data = {}

# -------- PARSE FILES --------
for uploaded_file in uploaded_files:
    month_year = os.path.splitext(uploaded_file.name)[0]
    df = pd.read_excel(uploaded_file, sheet_name="Total", header=None)

    records = []
    for portal, (c_task, c_val) in portals.items():
        portal_df = df[[c_task, c_val]].dropna(how="all")
        current_person = None

        for _, row in portal_df.iterrows():
            label = str(row.iloc[0]).strip()
            value = row.iloc[1]

            if label in known_tasks and current_person and pd.notna(value):
                records.append([month_year, portal, current_person, label, value])
            elif label and label not in known_tasks:
                current_person = label

    df_month = pd.DataFrame(
        records,
        columns=["Month", "Portal", "Person", "Task", "Completion"]
    )
    df_month["Completion"] = pd.to_numeric(df_month["Completion"], errors="coerce").fillna(0)

    monthly_data[month_year] = df_month
    all_data.append(df_month)

combined_df = pd.concat(all_data, ignore_index=True)

# -------- UI --------
st.title("📊 Task Completion Analyzer")
st.dataframe(combined_df, use_container_width=True)

st.divider()
st.markdown("## 📄 Generate Report")

if st.button("📥 Generate Full Word Report (.docx)"):
    with st.spinner("Generating report with embedded charts..."):
        report_path = generate_combined_word_report(combined_df, monthly_data)

        with open(report_path, "rb") as f:
            st.download_button(
                "⬇️ Download Report",
                f,
                file_name=os.path.basename(report_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
